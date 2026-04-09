"""
Stage 1: Question Paper Parsing (DOCX → Excel)

Reads a .docx question paper, extracts metadata and question-CO-marks mapping,
and writes a structured Excel output file.
"""

import sys
import json
import re
import os
from pathlib import Path
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
import pytesseract
from PIL import Image
from utils import validate_file_exists, find_value_after_keyword


tesseract_cmd = os.getenv("TESSERACT_CMD")
if tesseract_cmd:
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd


def read_docx(docx_path):
    """
    Read DOCX file and extract paragraphs and tables.

    Returns:
        (lines, tables): list of paragraph texts, list of table rows
    """
    doc = Document(docx_path)
    lines = []
    tables_data = []

    # Extract paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            tables_data.append(row_data)

    return lines, tables_data


def ocr_from_doc(docx_path):
    """
    Extract text from images embedded in DOCX using OCR.

    Returns:
        OCR text string
    """
    try:
        doc = Document(docx_path)
        ocr_text = ""

        # Iterate through document relationships to find images
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        # Extract image blob
                        image_part = rel.target_part
                        image_bytes = image_part.blob

                        # Save temporarily
                        temp_path = "temp_ocr.png"
                        with open(temp_path, "wb") as f:
                            f.write(image_bytes)

                        # OCR
                        img = Image.open(temp_path)
                        ocr_text += pytesseract.image_to_string(img) + "\n"

                        # Clean up
                        Path(temp_path).unlink(missing_ok=True)
                    except Exception:
                        pass
        except AttributeError:
            pass

        return ocr_text
    except Exception:
        return ""


def extract_metadata(lines, tables, docx_path):
    """
    Extract course metadata and CO information from DOCX content.

    Returns:
        (course_code, course_name, faculty, year, date, co_data)
    """
    full_text = "\n".join(lines) + "\n" + ocr_from_doc(docx_path)

    # Extract basic metadata
    course_code = find_value_after_keyword(tables, "Course Code")
    course_name = find_value_after_keyword(tables, "Course Name")
    faculty = find_value_after_keyword(tables, "Faculty")

    # Extract date (format: DD.MM.YYYY(...))
    date_match = re.search(r"\d{2}\.\d{2}\.\d{4}\(.*?\)", full_text)
    date = date_match.group() if date_match else "NA"

    # Extract academic year
    year_match = re.search(
        r"Academic year\s*[\d\-]+\s*\(.*?\)", full_text, re.IGNORECASE
    )
    year = year_match.group() if year_match else "NA"

    # Extract CO data from tables
    co_data = {}
    for row in tables:
        row_text = " ".join(row)

        co_match = re.search(r"(CO\d)", row_text)
        tsp_match = re.search(r"(TPS\d)", row_text)

        # Find last number in row (max marks)
        mark_matches = re.findall(r"\b\d+(?:\.\d+)?\b", row_text)
        mark = mark_matches[-1] if mark_matches else None

        if co_match:
            co_name = co_match.group()
            if co_name not in co_data:
                co_data[co_name] = {
                    "tsp": tsp_match.group() if tsp_match else "NA",
                    "mark": mark if mark else "NA",
                }

    return course_code, course_name, faculty, year, date, co_data


def process_docx(docx_path):
    """
    Main processing function for DOCX file.

    Returns:
        Pandas DataFrame with metadata and question structure
    """
    lines, tables = read_docx(docx_path)
    course_code, course_name, faculty, year, date, co_data = extract_metadata(
        lines, tables, docx_path
    )

    # Build metadata rows
    meta = [
        ["Course Code", course_code],
        ["Course Name", course_name],
        ["Faculty in-Charge", faculty],
        ["Academic Year", year],
        ["Date", date],
    ]

    for co in sorted(co_data.keys()):
        meta.append([co, co_data[co]["tsp"], co_data[co]["mark"]])

    # Convert metadata to DataFrame
    meta_df = pd.DataFrame(meta)

    # Extract questions from tables
    questions = []
    marks = []
    cos = []
    q_counter = 1
    part = "A"

    for row in tables:
        text = " ".join(row)

        # Detect part changes
        if "Part A" in text:
            part = "A"
            q_counter = 1
        elif "Part B" in text:
            part = "B"
            q_counter = 1
        elif "Part C" in text:
            part = "C"
            q_counter = 1

        m = None
        c = None

        # Extract marks and CO
        for cell in row:
            # Check if cell is numeric (marks)
            if re.fullmatch(r"\d+(?:\.\d+)?", cell.strip()):
                m = cell.strip()

            # Check if cell is a CO tag
            if re.match(r"CO\d", cell.strip(), re.IGNORECASE):
                c = cell.strip()

        # If both mark and CO found, add question
        if m and c:
            if part in ["A", "B"]:
                qno = f"{part}{q_counter}"
                q_counter += 1
            else:
                qno = text.split()[0] if text.split() else f"C{q_counter}"

            questions.append(qno)
            marks.append(m)
            cos.append(c)

    # Build question structure
    if questions:
        q_output = {i: [questions[i], marks[i], cos[i]] for i in range(len(questions))}
        q_df = pd.DataFrame.from_dict(q_output, orient="index").T
        final_df = pd.concat(
            [meta_df, pd.DataFrame([[""]]), q_df], ignore_index=True
        )
    else:
        final_df = meta_df

    return final_df


def main():
    try:
        # Parse arguments
        args = json.loads(sys.argv[1])
        docx_path = args["docx_path"]
        output_path = args["output_path"]

        # Validate input
        validate_file_exists(docx_path)

        # Process
        df = process_docx(docx_path)

        # Write to Excel
        df.to_excel(output_path, index=False, header=False)

        # Return success
        print(json.dumps({"status": "ok", "output_path": output_path}))

    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()
